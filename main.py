from flask import Flask, request, jsonify, send_file
import openai
import os
import configparser
import json
from docx import Document
import win32com.client as win32
import pythoncom

# Initialize Flask app
app = Flask(__name__)

# Set up the script directory and load the config file
script_directory = os.path.dirname(os.path.abspath(__file__))
config = configparser.ConfigParser()
config.read(os.path.join(script_directory, 'config.ini'))
api_key = config.get('OPENAI', 'api_key')
model = config.get('OPENAI', 'model')
temperature = config.getfloat('OPENAI', 'temperature')
max_tokens = config.getint('OPENAI', 'max_tokens')
top_p = config.getfloat('OPENAI', 'top_p')
frequency_penalty = config.getfloat('OPENAI', 'frequency_penalty')
presence_penalty = config.getfloat('OPENAI', 'presence_penalty')
finetuning_API_KEY = config.get('OPENAI','finetuning_API_KEY_4o')
finetuning_model = config.get('OPENAI','finetuning_model_4o')

# Directory Configuration
template_dir = os.path.join(script_directory,config.get('DIRECTORIES', 'template_dir'))
output_dir = os.path.join(script_directory,config.get('DIRECTORIES', 'output_dir'))

@app.route('/')
def home():
    return '''
    <html>
        <head>
            <title>Flask App Home</title>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    background-color: #f4f4f4;
                    margin: 50px;
                }
                h1 {
                    color: #333;
                }
                .container {
                    display: flex;
                    flex-direction: column;
                    align-items: center;
                }
                .btn {
                    padding: 10px 20px;
                    margin: 10px;
                    background-color: #007BFF;
                    color: white;
                    border: none;
                    cursor: pointer;
                    text-decoration: none;
                    font-size: 16px;
                }
                .btn:hover {
                    background-color: #0056b3;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Welcome to Orcas home page</h1>
                <p>Select an action:</p>
                <div>
                <button onclick="window.open('/download/deret_Leadbook.docx', '_blank')">Download your Leadbook</button>
                <button onclick="window.open('/download/deret_Brochure.docx', '_blank')">Download your company brochure</button>
                <button onclick="window.open('/download/Logistique_Brochure.docx', '_blank')">Download your Logistique offer brochure</button>
                <button onclick="window.open('/download/Transport_Brochure.docx', '_blank')">Download your Transport offer brochure</button>
            </div>
            </div>
        </body>
    </html>
    '''

@app.route('/generate_leadbook', methods=['POST'])
def generate_leadbook_file():
    data = request.json
    # Extract leadbook-related data from the input
    input_data = data.get('input')
    presentation_data = data.get('presentation_data')
    economic_data = data.get('economic_data')
    offers = data.get('offers')
    summary_text = data.get('summary_text')

    # Call the function that generates the leadbook
    leadbook_filename = generate_leadbook(input_data, presentation_data, economic_data, offers, summary_text)

    return jsonify({"leadbook": leadbook_filename}), 200

@app.route('/generate_brochure_company', methods=['POST'])
def generate_brochure_company():
    data = request.json
    # Extract company-related data from the input
    input = data.get('input')
    economic_data = data.get('economic_data')
    presentation_data = data.get('presentation_data')
    
    # Call the function that generates the brochure
    brochure_filename = generate_brochure(input,economic_data, None, False, presentation_data['Theme de la brochure'])
    
    return jsonify({"brochure": brochure_filename}), 200

@app.route('/generate_brochure_offer', methods=['POST'])
def generate_brochure_offer():
    data = request.json
    # Extract offer-related data from the input
    input = data.get('input')
    economic_data = data.get('economic_data')
    offer = data.get('offer')
    
    # Call the function that generates the brochure
    brochure_filename = generate_brochure(input,economic_data, offer,True, offer['Theme de la brochure'])
    
    return jsonify({"brochure": brochure_filename}), 200

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(output_dir, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"error": "File not found"}), 404
    
def ask_openai(api_key, model, system_content, prompt):
    # Query the OpenAI API with the given prompt
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_content},
            {"role": "user", "content": prompt}
        ],
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=top_p,
        frequency_penalty=frequency_penalty,
        presence_penalty=presence_penalty
    )
    return response.choices[0].message.content

def fine_tuning(finetuning_input):
    openai.api_key = finetuning_API_KEY
    completion = openai.chat.completions.create(
    model = finetuning_model,
    messages=[
        {"role": "system", "content": "Vous êtes un expert dans la génération de scripts personnalisés pour la prospection téléphonique, par email et sur LinkedIn, basés sur des données structurées et les outils de prospection marqués. Votre tâche consiste à créer des messages bien conçus, engageants et créatifs, adaptés aux besoins spécifiques et aux caractéristiques de chaque entreprise fournies dans les invites."},
        {"role": "user", "content": finetuning_input}
    ]
    )
    return(completion.choices[0].message.content)

def fetch_data(input, section, paragraph):
    # Determine system content and prompt based on whether a paragraph is required
    if paragraph:
        system_content = f"Tu es un assistant intelligent qui génère un leadbook pour le commercial de l'entreprise {input['Raison sociale']} ayant le SIRET {input['SIRET']}."
        prompt = (f"Rédige un paragraphe détaillé de la section '{section}' pour le commercial de l'entreprise {input['Raison sociale']} dont le secteur d'activité est {input['Secteur d\'activité']} sans inclure les titres. "
                  f"Consultez leur site web {input['Site web']} pour plus d'informations.")
    else:    
        system_content = f"Tu es un assistant intelligent qui aide à générer un leadbook pour le commercial de l'entreprise {input['Raison sociale']} ayant le SIRET {input['SIRET']} en répondant par un seul mot."
        prompt = f"{section} de l'entreprise {input['Raison sociale']} dont le secteur d'activité est {input['Secteur d\'activité']}. Consultez leur site web {input['Site web']} pour plus d'informations."
    
    return ask_openai(api_key, model, system_content, prompt)

def summerize_data(data):
    # Summarize the provided data into a detailed paragraph
    system_content = f"Tu es un assistant intelligent qui résume des sections du leadbook en un paragraphe détaillé contenant toutes les informations importantes."
    prompt = (f"Résume les données du dictionnaire {data} afin d\'obtenir un paragraphe détaillé. Je veux une réponse directe sans inclure des titres ou une introduction.Ne parlez pas de la brochure!")
    return ask_openai(api_key, model, system_content, prompt)          

def generate_presentation(input, presentation_data):
    # Generate the company presentation based on the input and selected options
    data = {}
    data['Présentation société'] = fetch_data(input, 'présentation société', True)

    if presentation_data['Groupe'] == 'Oui':
        data['Groupe'] = fetch_data(input, 'Quel est le groupe', False)
    
    if presentation_data['Certifications'] == 'Oui':
        data['Certifications'] = fetch_data(input, 'certifications', True)
    
    if presentation_data['Références'] == 'Oui':
        data['Références'] = fetch_data(input, 'références', True)
    
    if presentation_data['Année de création'] == 'Oui':
        data['Année de création'] = fetch_data(input, 'Quelle est l\'année de création', False)

    return summerize_data(data)

def generate_offer(input, offer):
    # Generate offer details based on the input and selected options
    data = dict(list(offer.items()))
    data['Description'] = fetch_data(input, f"offre {offer['Nom de l\'offre']}", True)
    
    return summerize_data(data)
    
def generate_brochure(input,economic_data, offer, brochure_offre, selected_theme):
    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    brochure_doc = None

    try:
        selected_theme = selected_theme.replace(' ', '_')
        brochure_template = os.path.join(template_dir, f"base_template_{selected_theme}.docx")

        try:
            brochure_doc = word.Documents.Open(brochure_template)
        except Exception as e:
            raise RuntimeError(f"Error opening the Word document: {e}")

        # Generate JSON content for the brochure
        system_content = f"Tu es un assistant intelligent qui génère une brochure sous forme JSON python pour l'entreprise {input['Raison sociale']}. Inclure des emojies convenables. Tu dois générer seulement le json sans commentaires et sans introduction."
        if brochure_offre:
            prompt = (f"Créez une brochure qui parle seulement de l'offre {offer['Nom de l\'offre']} proposé par la société {input['Raison sociale']} dont le secteur d'activité est {input['Secteur d\'activité']}. "
                        f"Vous pouvez consultez leur site web {input['Site web']} pour plus d'informations.Tu dois générer seulement le json sans commentaires et sans introduction.")
        else:
            prompt = (f"Créez une brochure pour la société {input['Raison sociale']} dont le secteur d'activité est {input['Secteur d\'activité']}. "
                        f"Vous pouvez consultez leur site web {input['Site web']} pour plus d'informations.Tu dois générer seulement le json sans commentaires et sans introduction.")
        prompt += "La réponse doit être un JSON string de cette forme: {\"company_name\":\"\",\"what_do_we_offer_section\":\"\",\"a_brief_description\":\"\",\"our_activity\":\"\",\"our_mission\":\"\",\"quote\":\"\",\"chiffre_affaires_en_M$\":\"\",\"phone_number\":\"\",\"email\":\"\",\"adresse\":\"\",\"site_web\":\"\",\"offer_service_1\":\"\",\"offer_service_2\":\"\", \"offer_service_3\":\"\",\"reason_to_trust_us_1\":\"\", \"reason_to_trust_us_2\":\"\", \"reason_to_trust_us_3\":\"\"}."

        response = ask_openai(api_key, model, system_content, prompt).strip()
        brochure_content = json.loads(response)
        
        placeholders = {
            '{company_name}': brochure_content['company_name'],
            '{what_do_we_offer}': brochure_content['what_do_we_offer_section'],
            '{a_brief_description}': brochure_content['a_brief_description'],
            '{our_activity}': brochure_content['our_activity'],
            '{our_mission}': brochure_content['our_mission'],
            '{quote}': brochure_content['quote'],
            '{chiffre_affaires}': economic_data['Chiffre d\'affaire en 2023'] if economic_data.get('Chiffre d\'affaire en 2023') else brochure_content['chiffre_affaires_en_M$'],
            '{tel}': brochure_content['phone_number'],
            '{email}': brochure_content['email'],
            '{adresse}': brochure_content['adresse'],
            '{site_web}': brochure_content['site_web'],
            '{service1}': brochure_content['offer_service_1'],
            '{service2}': brochure_content['offer_service_2'],
            '{service3}': brochure_content['offer_service_3'],
            '{reason1}': brochure_content['reason_to_trust_us_1'],
            '{reason2}': brochure_content['reason_to_trust_us_2'],
            '{reason3}': brochure_content['reason_to_trust_us_3']
        }

        # Ensure the document is not closed or corrupted before accessing shapes
        if brochure_doc is None:
            raise RuntimeError("The document could not be opened.")

        # Replace placeholders in the Word document with the generated content
        for shape in brochure_doc.Shapes:
            if shape.Type == 17 and shape.TextFrame.HasText:
                text = shape.TextFrame.TextRange.Text
                for placeholder, value in placeholders.items():
                    if placeholder in text:
                        shape.TextFrame.TextRange.Text = text.replace(placeholder, value)

        if brochure_offre:
            doc_filename = os.path.join(output_dir, f"{offer['Nom de l\'offre'].replace(' ', '_')}_Brochure.docx")
        else:
            doc_filename = os.path.join(output_dir, f"{input['Raison sociale'].replace(' ', '_')}_Brochure.docx")

        brochure_doc.SaveAs(doc_filename)
    finally:
        if brochure_doc is not None:
            brochure_doc.Close(SaveChanges=False)
        if word is not None:
            word.Quit()

    return doc_filename
    
def generate_leadbook(input, presentation_data, economic_data, offers, summary_text):
    # Generate the leadbook document based on the input data and summary
    nb_sections = 0
    summary = summary_text.split('\n')
    doc_filename = (os.path.join(output_dir,f"{input['Raison sociale'].replace(' ', '_')}_Leadbook.docx"))

    doc = Document()
    doc.add_heading(f"Informations de {input['Raison sociale']}", 0)
    for title in summary:
        if 'présentation de la société' in title.lower():
            nb_sections += 1
            doc.add_heading(f"{nb_sections}. Présentation de la société", 1)
            presentation_para = generate_presentation(input,economic_data, presentation_data)
            for line in presentation_para.split('\n'):
                doc.add_paragraph(line)
        
        if bool(economic_data) and 'données économiques' in title.lower():
            nb_sections += 1
            doc.add_heading(f"{nb_sections}. Données économiques", 1)
            economic_para = fetch_data(economic_data, "Données économiques", True)
            for line in economic_para.split('\n'):
                doc.add_paragraph(line)
        
        if bool(offers) and 'offres' in title.lower():
            nb_sections += 1
            doc.add_heading(f"{nb_sections}. Offres", 1)
            for index, offer in enumerate(offers):
                nb_script = 1
                doc.add_heading(f"{nb_sections}.{index+1} Offre {offer['Nom de l\'offre']}", 2)

                offer_para = generate_offer(input, economic_data, offer)

                for line in offer_para.split('\n'):
                    doc.add_paragraph(line)
                
                if offer['Outils de prospection']:
                    for outil in offer['Outils de prospection']:
                        doc.add_heading(f"{nb_sections}.{index+1}.{nb_script} Script de prospection {outil} pour l'offre {offer['Nom de l\'offre']}", 3)
                                    
                        finetuning_input = {
                            'Outils de prospection': outil,
                            'Raison sociale': input['Raison sociale'],
                            'SIRET': input['SIRET'],
                            'Secteur d\'activité': input['Secteur d\'activité'],
                            'Site web': input['Site web'],
                            'Nom de l\'offre': offer['Nom de l\'offre'],
                            'Référence 1': offer.get('Référence 1'),
                            'Référence 2': offer.get('Référence 2'),
                            'Référence 3': offer.get('Référence 3'),
                            'Part du CA': offer.get('Part du CA'),
                            'Part du marché': offer.get('Part du marché'),
                            'Pays de livraison': offer.get('Pays de livraison'),
                            'Offre stratégique': offer.get('Offre stratégique'),
                            'Marché cible': offer.get('Marché cible'),
                            'Partenariats': offer.get('Partenariats')
                        }

                        # Generate fine-tuned script
                        finetuned_script = fine_tuning(f"{finetuning_input}")
                        doc.add_paragraph(finetuned_script)
                        nb_script += 1

    doc.save(doc_filename)
    return doc_filename

# Running the Flask app
if __name__ == '__main__':
    app.run(debug=True)